"""
文檔處理管道
支援多種文件格式的解析、清理和分塊
"""

import asyncio
import json
import mimetypes
import re
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union
import hashlib
from datetime import datetime

# 文檔解析庫
try:
    import pypdf2
except ImportError:
    try:
        from pypdf import PdfReader as pypdf2_PdfReader
        # 為兼容性創建一個包裝類
        class pypdf2:
            PdfReader = pypdf2_PdfReader
    except ImportError:
        pypdf2 = None
import docx
from bs4 import BeautifulSoup
import pandas as pd
from markdown import markdown
import ast

from config.settings import get_settings


class DocumentMetadata:
    """文檔元數據"""
    
    def __init__(self, file_path: Union[str, Path]):
        self.file_path = Path(file_path)
        self.file_name = self.file_path.name
        self.file_size = self.file_path.stat().st_size if self.file_path.exists() else 0
        self.file_type = self.file_path.suffix.lower()
        self.mime_type = mimetypes.guess_type(str(self.file_path))[0]
        self.created_at = datetime.now()
        self.modified_at = datetime.fromtimestamp(self.file_path.stat().st_mtime) if self.file_path.exists() else None
        self.hash = self._calculate_hash()
    
    def _calculate_hash(self) -> str:
        """計算文件哈希值"""
        if not self.file_path.exists():
            return ""
        
        hash_sha256 = hashlib.sha256()
        with open(self.file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def to_dict(self) -> Dict[str, Any]:
        """轉換為字典"""
        return {
            "file_path": str(self.file_path),
            "file_name": self.file_name,
            "file_size": self.file_size,
            "file_type": self.file_type,
            "mime_type": self.mime_type,
            "created_at": self.created_at.isoformat(),
            "modified_at": self.modified_at.isoformat() if self.modified_at else None,
            "hash": self.hash
        }


class DocumentChunk:
    """文檔塊"""
    
    def __init__(self, content: str, metadata: DocumentMetadata, 
                 chunk_index: int = 0, chunk_type: str = "text",
                 start_char: int = 0, end_char: int = 0):
        self.content = content
        self.metadata = metadata
        self.chunk_index = chunk_index
        self.chunk_type = chunk_type
        self.start_char = start_char
        self.end_char = end_char
        self.word_count = len(content.split())
        self.char_count = len(content)
        self.hash = hashlib.sha256(content.encode()).hexdigest()
    
    def to_dict(self) -> Dict[str, Any]:
        """轉換為字典"""
        return {
            "content": self.content,
            "metadata": self.metadata.to_dict(),
            "chunk_index": self.chunk_index,
            "chunk_type": self.chunk_type,
            "start_char": self.start_char,
            "end_char": self.end_char,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "hash": self.hash
        }


class BaseDocumentParser(ABC):
    """文檔解析器基類"""
    
    @abstractmethod
    def can_parse(self, file_path: Path) -> bool:
        """檢查是否能解析該文件"""
        pass
    
    @abstractmethod
    async def parse(self, file_path: Path) -> str:
        """解析文檔內容"""
        pass
    
    def clean_text(self, text: str) -> str:
        """清理文本內容"""
        # 移除多餘空白字符
        text = re.sub(r'\s+', ' ', text)
        # 移除特殊字符
        text = re.sub(r'[^\w\s\u4e00-\u9fff.,!?;:()"\'-]', '', text)
        return text.strip()


class PDFParser(BaseDocumentParser):
    """PDF文檔解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.pdf'
    
    async def parse(self, file_path: Path) -> str:
        """解析PDF文檔"""
        if pypdf2 is None:
            raise ValueError("PDF parsing library not available. Please install pypdf or pypdf2.")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf2.PdfReader(file)
                text_content = []
                
                # 兼容不同版本的API
                if hasattr(pdf_reader, 'pages'):
                    pages = pdf_reader.pages
                elif hasattr(pdf_reader, 'numPages'):
                    pages = [pdf_reader.getPage(i) for i in range(pdf_reader.numPages)]
                else:
                    raise ValueError("Unsupported PDF library version")
                
                for page in pages:
                    if hasattr(page, 'extract_text'):
                        text_content.append(page.extract_text())
                    elif hasattr(page, 'extractText'):
                        text_content.append(page.extractText())
                    else:
                        text_content.append("")
                
                return self.clean_text('\n'.join(text_content))
        except Exception as e:
            raise ValueError(f"Failed to parse PDF file {file_path}: {str(e)}")


class WordParser(BaseDocumentParser):
    """Word文檔解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.docx', '.doc']
    
    async def parse(self, file_path: Path) -> str:
        """解析Word文檔"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # 處理表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            return self.clean_text('\n'.join(text_content))
        except Exception as e:
            raise ValueError(f"Failed to parse Word file {file_path}: {str(e)}")


class TextParser(BaseDocumentParser):
    """純文本解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.txt', '.md', '.rst']
    
    async def parse(self, file_path: Path) -> str:
        """解析文本文檔"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # 如果是Markdown，轉換為純文本
            if file_path.suffix.lower() == '.md':
                html = markdown(content)
                soup = BeautifulSoup(html, 'html.parser')
                content = soup.get_text()
            
            return self.clean_text(content)
        except Exception as e:
            raise ValueError(f"Failed to parse text file {file_path}: {str(e)}")


class PythonParser(BaseDocumentParser):
    """Python代碼解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.py'
    
    async def parse(self, file_path: Path) -> str:
        """解析Python代碼"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            # 解析AST以提取文檔字符串和註釋
            try:
                tree = ast.parse(content)
                docstrings = []
                
                for node in ast.walk(tree):
                    if isinstance(node, (ast.FunctionDef, ast.ClassDef, ast.Module)):
                        if (isinstance(node.body[0], ast.Expr) and 
                            isinstance(node.body[0].value, ast.Str)):
                            docstrings.append(node.body[0].value.s)
                
                # 提取註釋
                comments = re.findall(r'#.*', content)
                
                # 組合代碼結構信息
                structured_content = f"Code file: {file_path.name}\n"
                structured_content += f"Content:\n{content}\n"
                if docstrings:
                    structured_content += f"Docstrings:\n" + '\n'.join(docstrings) + "\n"
                if comments:
                    structured_content += f"Comments:\n" + '\n'.join(comments)
                
                return structured_content
                
            except SyntaxError:
                # 如果解析失敗，返回原始內容
                return f"Code file: {file_path.name}\nContent:\n{content}"
                
        except Exception as e:
            raise ValueError(f"Failed to parse Python file {file_path}: {str(e)}")


class HTMLParser(BaseDocumentParser):
    """HTML文檔解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in ['.html', '.htm']
    
    async def parse(self, file_path: Path) -> str:
        """解析HTML文檔"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # 移除腳本和樣式標籤
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text()
            return self.clean_text(text)
            
        except Exception as e:
            raise ValueError(f"Failed to parse HTML file {file_path}: {str(e)}")


class JSONParser(BaseDocumentParser):
    """JSON文檔解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.json'
    
    async def parse(self, file_path: Path) -> str:
        """解析JSON文檔"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                data = json.load(file)
            
            # 將JSON轉換為可讀文本
            return self._json_to_text(data, file_path.name)
            
        except Exception as e:
            raise ValueError(f"Failed to parse JSON file {file_path}: {str(e)}")
    
    def _json_to_text(self, data: Any, filename: str) -> str:
        """將JSON數據轉換為文本描述"""
        content = [f"JSON file: {filename}"]
        
        def describe_value(value, path=""):
            if isinstance(value, dict):
                for key, val in value.items():
                    new_path = f"{path}.{key}" if path else key
                    content.append(f"{new_path}: {type(val).__name__}")
                    if isinstance(val, (str, int, float, bool)):
                        content.append(f"  Value: {val}")
                    else:
                        describe_value(val, new_path)
            elif isinstance(value, list):
                content.append(f"{path}: List with {len(value)} items")
                for i, item in enumerate(value[:5]):  # 只處理前5個項目
                    describe_value(item, f"{path}[{i}]")
            else:
                content.append(f"{path}: {value}")
        
        describe_value(data)
        return '\n'.join(content)


class CSVParser(BaseDocumentParser):
    """CSV文檔解析器"""
    
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == '.csv'
    
    async def parse(self, file_path: Path) -> str:
        """解析CSV文檔"""
        try:
            df = pd.read_csv(file_path)
            
            content = [f"CSV file: {file_path.name}"]
            content.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
            content.append(f"Columns: {', '.join(df.columns.tolist())}")
            
            # 添加數據樣本
            content.append("Sample data:")
            content.append(df.head().to_string())
            
            # 添加數據描述
            if df.select_dtypes(include=['number']).columns.any():
                content.append("Numerical summary:")
                content.append(df.describe().to_string())
            
            return '\n'.join(content)
            
        except Exception as e:
            raise ValueError(f"Failed to parse CSV file {file_path}: {str(e)}")


class TextChunker:
    """文本分塊器"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def chunk_text(self, text: str, metadata: DocumentMetadata) -> List[DocumentChunk]:
        """將文本分塊"""
        if len(text) <= self.chunk_size:
            return [DocumentChunk(text, metadata, 0, "text", 0, len(text))]
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            
            # 嘗試在句號或換行符處分割
            if end < len(text):
                last_sentence = text.rfind('.', start, end)
                last_newline = text.rfind('\n', start, end)
                
                split_pos = max(last_sentence, last_newline)
                if split_pos > start + self.chunk_size // 2:
                    end = split_pos + 1
            
            chunk_content = text[start:end].strip()
            if chunk_content:
                chunks.append(DocumentChunk(
                    content=chunk_content,
                    metadata=metadata,
                    chunk_index=chunk_index,
                    chunk_type="text",
                    start_char=start,
                    end_char=end
                ))
                chunk_index += 1
            
            # 計算下一個塊的起始位置（考慮重疊）
            start = max(end - self.chunk_overlap, start + 1)
        
        return chunks


class DocumentProcessingPipeline:
    """文檔處理管道"""
    
    def __init__(self):
        self.settings = get_settings()
        self.parsers = [
            PDFParser(),
            WordParser(),
            TextParser(),
            PythonParser(),
            HTMLParser(),
            JSONParser(),
            CSVParser()
        ]
        
        self.chunker = TextChunker(
            chunk_size=self.settings.rag.chunk_size,
            chunk_overlap=self.settings.rag.chunk_overlap
        )
    
    def _get_parser(self, file_path: Path) -> Optional[BaseDocumentParser]:
        """獲取適合的解析器"""
        for parser in self.parsers:
            if parser.can_parse(file_path):
                return parser
        return None
    
    async def process_file(self, file_path: Union[str, Path]) -> List[DocumentChunk]:
        """處理單個文件"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.stat().st_size > self.settings.storage.max_file_size:
            raise ValueError(f"File too large: {file_path}")
        
        if file_path.suffix.lower() not in self.settings.storage.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        parser = self._get_parser(file_path)
        if not parser:
            raise ValueError(f"No parser available for file: {file_path}")
        
        try:
            # 解析文檔內容
            content = await parser.parse(file_path)
            
            # 創建元數據
            metadata = DocumentMetadata(file_path)
            
            # 分塊處理
            chunks = self.chunker.chunk_text(content, metadata)
            
            return chunks
            
        except Exception as e:
            raise ValueError(f"Failed to process file {file_path}: {str(e)}")
    
    async def process_directory(self, directory_path: Union[str, Path], 
                              recursive: bool = True) -> List[DocumentChunk]:
        """批量處理目錄中的文件"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            raise ValueError(f"Invalid directory: {directory_path}")
        
        # 收集所有支援的文件
        files = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory_path.glob(pattern):
            if (file_path.is_file() and 
                file_path.suffix.lower() in self.settings.storage.supported_formats):
                files.append(file_path)
        
        # 並行處理文件
        all_chunks = []
        semaphore = asyncio.Semaphore(self.settings.agent.max_parallel_tasks)
        
        async def process_with_semaphore(file_path):
            async with semaphore:
                try:
                    return await self.process_file(file_path)
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
                    return []
        
        # 執行並行處理
        tasks = [process_with_semaphore(file_path) for file_path in files]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # 收集結果
        for result in results:
            if isinstance(result, list):
                all_chunks.extend(result)
        
        return all_chunks
    
    async def process_batch(self, file_paths: List[Union[str, Path]]) -> List[DocumentChunk]:
        """批量處理文件列表"""
        all_chunks = []
        semaphore = asyncio.Semaphore(self.settings.agent.max_parallel_tasks)
        
        async def process_with_semaphore(file_path):
            async with semaphore:
                try:
                    return await self.process_file(file_path)
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
                    return []
        
        tasks = [process_with_semaphore(file_path) for file_path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        for result in results:
            if isinstance(result, list):
                all_chunks.extend(result)
        
        return all_chunks
